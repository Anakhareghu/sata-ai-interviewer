"""
Resume Parser Service

Extracts information from PDF and DOCX resumes using:
- pdfplumber for PDF text extraction
- python-docx for DOCX parsing
- Regex patterns for skill, project, and experience extraction
"""

import os
import re
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parse and extract information from resumes."""
    
    def __init__(self):
        # Comprehensive skill lists
        self.technical_skills = {
            "programming": ["python", "java", "javascript", "typescript", "c++", "c#", "ruby", "go", "rust", "php", "swift", "kotlin"],
            "web": ["html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask", "fastapi", "spring", "nextjs"],
            "database": ["sql", "mysql", "postgresql", "mongodb", "redis", "sqlite", "oracle", "dynamodb", "firebase"],
            "cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "ci/cd"],
            "ml": ["machine learning", "deep learning", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "nlp", "computer vision"],
            "tools": ["git", "linux", "bash", "jira", "figma", "postman", "webpack", "vite"]
        }
        
        self.soft_skills = [
            "leadership", "communication", "teamwork", "problem solving", "analytical",
            "creative", "adaptable", "organized", "detail-oriented", "proactive"
        ]
    
    async def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a resume file and extract information.
        
        Args:
            file_path: Path to resume file (PDF or DOCX)
            
        Returns:
            Dictionary with extracted information
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        
        if ext == ".pdf":
            text = await self._extract_pdf_text(file_path)
        elif ext in [".docx", ".doc"]:
            text = await self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        logger.info(f"📄 Extracted {len(text)} characters from resume")
        
        return await self._analyze_text(text)
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF."""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    async def _analyze_text(self, text: str) -> Dict[str, Any]:
        """Analyze resume text and extract information."""
        text_lower = text.lower()
        
        return {
            "raw_text": text[:5000],  # Store first 5000 chars
            "extracted_skills": self._extract_skills(text_lower),
            "projects": self._extract_projects(text),
            "work_experience": self._extract_experience(text),
            "education": self._extract_education(text),
            "contact_info": self._extract_contact_info(text),
            "domain_specialization": self._determine_domain(text_lower),
            "experience_level": self._determine_experience_level(text_lower)
        }
    
    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract technical and soft skills."""
        technical = []
        soft = []
        
        # Extract technical skills
        for category, skills in self.technical_skills.items():
            for skill in skills:
                if skill in text:
                    technical.append(skill.title())
        
        # Extract soft skills
        for skill in self.soft_skills:
            if skill in text:
                soft.append(skill.title())
        
        return {
            "technical": list(set(technical)),
            "soft": list(set(soft))
        }
    
    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract project information."""
        projects = []
        
        # Look for project patterns
        project_patterns = [
            r"(?:project|built|developed|created|designed)[\s:]+([A-Z][^\n.]{10,100})",
            r"([A-Z][a-zA-Z\s]+)[\s]*[-–][\s]*(?:A|An|The)[^\n]{20,150}"
        ]
        
        for pattern in project_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches[:5]:  # Limit to 5 projects
                if len(match.strip()) > 10:
                    projects.append({
                        "name": match.strip()[:100],
                        "description": ""
                    })
        
        return projects[:5]
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience."""
        experience = []
        
        # Look for job title patterns
        job_patterns = [
            r"((?:Software|Web|Full[ -]?Stack|Frontend|Backend|Data|ML|AI)\s*(?:Engineer|Developer|Intern|Analyst))",
            r"((?:Junior|Senior|Lead)\s*(?:Developer|Engineer))",
            r"(Intern(?:ship)?(?:\s+at)?)"
        ]
        
        for pattern in job_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches[:3]:
                experience.append({
                    "title": match.strip(),
                    "company": "",
                    "duration": ""
                })
        
        return experience
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information."""
        education = []
        
        # Degree patterns
        degree_patterns = [
            r"(B\.?(?:Tech|E|S|Sc|A)?\.?\s*(?:in)?\s*(?:Computer|Information|Software|Electronics|Mechanical)?[^\n]{0,50})",
            r"(M\.?(?:Tech|S|Sc|A)?\.?\s*(?:in)?\s*(?:Computer|Information|Software)?[^\n]{0,50})",
            r"(Bachelor|Master)(?:'s)?\s+(?:of|in)\s+[^\n]{10,80}"
        ]
        
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches[:2]:
                education.append({
                    "degree": match.strip()[:100],
                    "institution": "",
                    "year": ""
                })
        
        return education
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information."""
        contact = {}
        
        # Email
        email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
        emails = re.findall(email_pattern, text)
        if emails:
            contact["email"] = emails[0]
        
        # Phone
        phone_pattern = r"(?:\+\d{1,3}[\s-]?)?\d{10}|\d{3}[\s-]\d{3}[\s-]\d{4}"
        phones = re.findall(phone_pattern, text)
        if phones:
            contact["phone"] = phones[0]
        
        # LinkedIn
        linkedin_pattern = r"linkedin\.com/in/([a-zA-Z0-9_-]+)"
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact["linkedin"] = f"linkedin.com/in/{linkedin[0]}"
        
        # GitHub
        github_pattern = r"github\.com/([a-zA-Z0-9_-]+)"
        github = re.findall(github_pattern, text, re.IGNORECASE)
        if github:
            contact["github"] = f"github.com/{github[0]}"
        
        return contact
    
    def _determine_domain(self, text: str) -> str:
        """Determine the candidate's domain specialization."""
        domain_keywords = {
            "AI/ML": ["machine learning", "deep learning", "neural network", "tensorflow", "pytorch", "nlp", "computer vision"],
            "Web Development": ["react", "angular", "vue", "node.js", "frontend", "backend", "fullstack", "web development"],
            "Data Science": ["data science", "data analysis", "pandas", "visualization", "statistics", "data engineer"],
            "Mobile Development": ["android", "ios", "flutter", "react native", "mobile app"],
            "DevOps": ["devops", "kubernetes", "docker", "ci/cd", "aws", "azure", "cloud"],
            "Cybersecurity": ["security", "penetration", "vulnerability", "encryption", "cybersecurity"]
        }
        
        scores = {}
        for domain, keywords in domain_keywords.items():
            score = sum(1 for kw in keywords if kw in text)
            if score > 0:
                scores[domain] = score
        
        if scores:
            return max(scores, key=scores.get)
        return "General Technology"
    
    def _determine_experience_level(self, text: str) -> str:
        """Determine experience level."""
        if any(word in text for word in ["senior", "lead", "principal", "architect", "5+ years", "7+ years"]):
            return "Senior"
        elif any(word in text for word in ["junior", "entry", "fresher", "graduate", "intern"]):
            return "Entry Level"
        else:
            return "Mid Level"


# Global instance
resume_parser = ResumeParser()
